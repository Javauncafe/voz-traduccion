
import streamlit as st
import speech_recognition as sr
from googletrans import Translator
from docx import Document
from fpdf import FPDF
import tempfile

st.set_page_config(page_title="Voz a Texto y Traducción", layout="centered")

st.title("🗣️ Transcripción y Traducción de Voz (Inglés a Español)")

recognizer = sr.Recognizer()
translator = Translator()

def transcribe_audio():
    with sr.Microphone() as source:
        st.info("Habla en inglés ahora...")
        audio = recognizer.listen(source, timeout=5)
        st.success("Grabación completada.")
        try:
            text = recognizer.recognize_google(audio, language="en-US")
            return text
        except sr.UnknownValueError:
            return "No se pudo entender el audio."
        except sr.RequestError:
            return "Error de conexión con el servicio de reconocimiento."

if st.button("🎙️ Grabar Voz en Inglés"):
    english_text = transcribe_audio()
    st.subheader("Texto en Inglés:")
    st.write(english_text)

    with st.spinner("Traduciendo al español..."):
        translated_text = translator.translate(english_text, src='en', dest='es').text
        st.subheader("Texto en Español:")
        st.write(translated_text)

        st.session_state['english'] = english_text
        st.session_state['spanish'] = translated_text

def generar_docx(ingles, espanol):
    doc = Document()
    doc.add_heading("Transcripción y Traducción", level=1)
    doc.add_heading("Inglés:", level=2)
    doc.add_paragraph(ingles)
    doc.add_heading("Español:", level=2)
    doc.add_paragraph(espanol)
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    return temp_file.name

def generar_pdf(ingles, espanol):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Transcripción y Traducción", ln=True, align='C')
    pdf.ln(10)
    pdf.cell(200, 10, txt="Inglés:", ln=True)
    pdf.multi_cell(0, 10, txt=ingles)
    pdf.ln(5)
    pdf.cell(200, 10, txt="Español:", ln=True)
    pdf.multi_cell(0, 10, txt=espanol)
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    pdf.output(temp_file.name)
    return temp_file.name

if 'english' in st.session_state and 'spanish' in st.session_state:
    st.markdown("### 📥 Descargar Resultado:")
    col1, col2 = st.columns(2)
    with col1:
        docx_file = generar_docx(st.session_state['english'], st.session_state['spanish'])
        with open(docx_file, "rb") as f:
            st.download_button("⬇️ Descargar Word", f, file_name="transcripcion.docx")
    with col2:
        pdf_file = generar_pdf(st.session_state['english'], st.session_state['spanish'])
        with open(pdf_file, "rb") as f:
            st.download_button("⬇️ Descargar PDF", f, file_name="transcripcion.pdf")
